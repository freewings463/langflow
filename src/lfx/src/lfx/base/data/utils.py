"""
模块名称：utils

本模块提供文本/结构化文件解析、编码检测、目录遍历与并发加载工具。
主要功能包括：
- 功能1：解析 JSON/YAML/XML 与通用文本归一化。
- 功能2：本地与对象存储的文本、PDF、DOCX 读取。
- 功能3：目录遍历筛选与并发加载 `Data`。

使用场景：文件类组件需要统一读取与解析逻辑时。
关键组件：
- 函数 `parse_structured_text`
- 函数 `retrieve_file_paths`
- 函数 `parse_text_file_to_data` / `parse_text_file_to_data_async`
- 函数 `parallel_load_data`

设计背景：统一文件读取与解析行为，避免组件各自实现造成差异。
注意事项：DOCX 在对象存储场景需要临时文件；编码检测可能误判。
"""

import contextlib
import tempfile
import unicodedata
from collections.abc import Callable
from concurrent import futures
from io import BytesIO
from pathlib import Path

import chardet
import orjson
import yaml
from defusedxml import ElementTree
from pypdf import PdfReader

from lfx.base.data.storage_utils import read_file_bytes
from lfx.schema.data import Data
from lfx.services.deps import get_settings_service
from lfx.utils.async_helpers import run_until_complete

# 注意：可直接 `file.read()` 读取且通常为纯文本的文件类型。
TEXT_FILE_TYPES = [
    "csv",
    "json",
    "pdf",
    "txt",
    "md",
    "mdx",
    "yaml",
    "yml",
    "xml",
    "html",
    "htm",
    "docx",
    "py",
    "sh",
    "sql",
    "js",
    "ts",
    "tsx",
]

IMG_FILE_TYPES = ["jpg", "jpeg", "png", "bmp", "image"]


def parse_structured_text(text: str, file_path: str) -> str | dict | list:
    """按文件扩展名解析 JSON/YAML/XML 并返回规范化结果。

    契约：JSON 返回字符串化结果；YAML 返回对象；XML 返回字符串。
    关键路径：识别扩展名 -> 解析 -> 规范化输出。
    决策：
    问题：结构化文本需要按格式解析，同时保持文本一致性。
    方案：依据扩展名选择解析器并对 JSON 内部文本做归一化。
    代价：解析失败会抛异常，需调用方处理。
    重评：当引入统一解析框架或 schema 校验时。
    """
    if file_path.endswith(".json"):
        loaded_json = orjson.loads(text)
        if isinstance(loaded_json, dict):
            loaded_json = {k: normalize_text(v) if isinstance(v, str) else v for k, v in loaded_json.items()}
        elif isinstance(loaded_json, list):
            loaded_json = [normalize_text(item) if isinstance(item, str) else item for item in loaded_json]
        return orjson.dumps(loaded_json).decode("utf-8")

    if file_path.endswith((".yaml", ".yml")):
        return yaml.safe_load(text)

    if file_path.endswith(".xml"):
        xml_element = ElementTree.fromstring(text)
        return ElementTree.tostring(xml_element, encoding="unicode")

    return text


def normalize_text(text):
    """进行 Unicode 规范化（NFKD）。

    契约：输入字符串，返回规范化字符串。
    关键路径：调用 `unicodedata.normalize("NFKD", ...)`。
    决策：
    问题：不同来源文本可能包含兼容字符或组合字符。
    方案：统一使用 NFKD 做兼容分解。
    代价：部分字符可能被展开，显示与原文略有差异。
    重评：当需要保留原始字形时。
    """
    return unicodedata.normalize("NFKD", text)


def is_hidden(path: Path) -> bool:
    """判断路径是否为隐藏文件（以 `.` 开头）。

    契约：仅检查文件名首字符，不解析系统属性。
    关键路径：读取 `path.name` 并检查前缀。
    决策：
    问题：跨平台隐藏文件标准不一致。
    方案：以 Unix 约定的前缀 `.` 作为判断依据。
    代价：Windows 隐藏属性不会被识别。
    重评：当需要跨平台隐藏检测时。
    """
    return path.name.startswith(".")


def format_directory_path(path: str) -> str:
    """格式化目录路径以避免换行注入。

    契约：仅替换换行符为可见 `\\n`。
    关键路径：执行字符串替换并返回新路径。
    决策：
    问题：路径中换行会破坏日志与错误消息展示。
    方案：将 `\n` 转义为字面量。
    代价：仅处理换行，未处理其他控制字符。
    重评：当需要更严格的路径清洗规则时。
    """
    return path.replace("\n", "\\n")


# 迁移背景：DirectoryComponent(1.0.19) 仍使用位置参数调用，保留 `FBT001` 例外。
def retrieve_file_paths(
    path: str,
    load_hidden: bool,  # noqa: FBT001
    recursive: bool,  # noqa: FBT001
    depth: int,
    types: list[str] = TEXT_FILE_TYPES,
) -> list[str]:
    """遍历目录并返回匹配类型的文件路径列表。

    契约：`path` 必须存在且为目录；`types` 为空表示不过滤后缀。
    关键路径：
    1) 校验目录有效性；
    2) 根据 `recursive/depth` 生成遍历器；
    3) 按后缀与隐藏文件规则过滤。
    异常流：路径不存在或非目录时抛 `ValueError`。
    决策：
    问题：不同组件对隐藏文件与递归深度有差异需求。
    方案：提供 `load_hidden/recursive/depth` 参数组合。
    代价：深度递归可能带来较高 I/O 成本。
    重评：当目录扫描交由专用索引服务时。
    """
    path = format_directory_path(path)
    path_obj = Path(path)
    if not path_obj.exists() or not path_obj.is_dir():
        msg = f"Path {path} must exist and be a directory."
        raise ValueError(msg)

    def match_types(p: Path) -> bool:
        return any(p.suffix == f".{t}" for t in types) if types else True

    def is_not_hidden(p: Path) -> bool:
        return not is_hidden(p) or load_hidden

    def walk_level(directory: Path, max_depth: int):
        directory = directory.resolve()
        prefix_length = len(directory.parts)
        for p in directory.rglob("*" if recursive else "[!.]*"):
            if len(p.parts) - prefix_length <= max_depth:
                yield p

    glob = "**/*" if recursive else "*"
    paths = walk_level(path_obj, depth) if depth else path_obj.glob(glob)
    return [str(p) for p in paths if p.is_file() and match_types(p) and is_not_hidden(p)]


def partition_file_to_data(file_path: str, *, silent_errors: bool) -> Data | None:
    """使用 `unstructured` 的 `partition` 解析文件并生成 `Data`。

    契约：解析成功返回 `Data`，失败在 `silent_errors=True` 时返回 `None`。
    关键路径：调用 `partition` -> 拼接元素文本 -> 组装 `Data`。
    副作用：依赖 `unstructured`，可能触发外部模型或解析器加载。
    决策：
    问题：通用分区解析能力不稳定且依赖较重。
    方案：提供可选路径，失败时可静默。
    代价：解析结果可能与文件类型专用解析存在差异。
    重评：当标准解析链路稳定后再启用默认路径。
    """
    # 注意：`partition` 可能触发较重的依赖加载。
    from unstructured.partition.auto import partition

    try:
        elements = partition(file_path)
    except Exception as e:
        if not silent_errors:
            msg = f"Error loading file {file_path}: {e}"
            raise ValueError(msg) from e
        return None

    # 实现：组装 `Data`，附带 `file_path` 元信息。
    text = "\n\n".join([str(el) for el in elements])
    metadata = elements.metadata if hasattr(elements, "metadata") else {}
    metadata["file_path"] = file_path
    return Data(text=text, data=metadata)


def read_text_file(file_path: str) -> str:
    """读取本地文本文件并自动检测编码。

    契约：仅支持本地路径；返回解码后的字符串。
    关键路径：读取字节 -> `chardet` 检测 -> 解码。
    决策：
    问题：不同来源文件编码不一致，直接读取易乱码。
    方案：使用 `chardet` 检测并对常见误判编码回退到 `utf-8`。
    代价：检测过程有额外开销，且可能误判。
    重评：当编码已由上游统一或可外部传入时。
    """
    file_path_ = Path(file_path)
    raw_data = file_path_.read_bytes()
    result = chardet.detect(raw_data)
    encoding = result["encoding"]

    if encoding in {"Windows-1252", "Windows-1254", "MacRoman"}:
        encoding = "utf-8"

    return file_path_.read_text(encoding=encoding)


async def read_text_file_async(file_path: str) -> str:
    """读取文本并自动检测编码（异步，支持对象存储）。

    契约：`file_path` 可为 `flow_id/filename` 或本地路径；返回字符串。
    关键路径：读取字节 -> 检测编码 -> 解码。
    决策：
    问题：对象存储无法直接走本地文件读取。
    方案：使用 `read_file_bytes` 获取字节后解码。
    代价：对大文件会增加内存占用。
    重评：当支持流式解码时。
    """
    from .storage_utils import read_file_bytes

    # 实现：统一使用存储感知读取获取字节内容。
    raw_data = await read_file_bytes(file_path)

    # 实现：自动检测编码，兼容多来源文本。
    result = chardet.detect(raw_data)
    encoding = result.get("encoding")

    # 注意：检测失败或误判时回退到 `utf-8`，避免直接抛异常。
    if not encoding or encoding in {"Windows-1252", "Windows-1254", "MacRoman"}:
        encoding = "utf-8"

    return raw_data.decode(encoding, errors="replace")


def read_docx_file(file_path: str) -> str:
    """读取本地 DOCX 并抽取文本。

    契约：仅支持本地路径；返回段落以 `\\n\\n` 拼接的文本。
    关键路径：加载 `Document` -> 遍历段落 -> 拼接文本。
    决策：
    问题：`python-docx` 仅接受文件路径，无法直接解析字节流。
    方案：本函数限定本地路径，S3 场景使用 `read_docx_file_async`。
    代价：需要区分存储类型。
    重评：当 `python-docx` 支持流式读取时。
    """
    from docx import Document

    doc = Document(file_path)
    return "\n\n".join([p.text for p in doc.paragraphs])


async def read_docx_file_async(file_path: str) -> str:
    """读取 DOCX 并抽取文本（异步，支持对象存储）。

    契约：`s3` 场景会下载到临时文件；返回段落文本拼接结果。
    关键路径：读取字节 -> 写入临时文件 -> 解析段落。
    决策：
    问题：`python-docx` 仅支持文件路径。
    方案：对象存储场景写入临时文件再解析。
    代价：产生临时文件，需要显式清理。
    重评：当 `python-docx` 支持 `BytesIO` 时。
    """
    from docx import Document

    from .storage_utils import read_file_bytes

    settings = get_settings_service().settings

    if settings.storage_type == "local":
        # 实现：本地存储直接读取，避免临时文件。
        doc = Document(file_path)
        return "\n\n".join([p.text for p in doc.paragraphs])

    # 注意：`s3` 需临时文件，`python-docx` 不支持 `BytesIO`。
    content = await read_file_bytes(file_path)

    # 实现：使用原始后缀创建临时文件，确保解析器识别。
    suffix = Path(file_path.split("/")[-1]).suffix
    with tempfile.NamedTemporaryFile(mode="wb", suffix=suffix, delete=False) as tmp_file:
        tmp_file.write(content)
        temp_path = tmp_file.name

    try:
        doc = Document(temp_path)
        return "\n\n".join([p.text for p in doc.paragraphs])
    finally:
        with contextlib.suppress(Exception):
            Path(temp_path).unlink()


def parse_pdf_to_text(file_path: str) -> str:
    """读取本地 PDF 并抽取所有页面文本。

    契约：仅支持本地路径；返回按页面拼接的字符串。
    关键路径：打开文件 -> `PdfReader` 解析 -> 汇总页面文本。
    决策：
    问题：PDF 解析需要二进制读取。
    方案：使用 `pypdf.PdfReader` 逐页抽取。
    代价：扫描型 PDF 可能返回空文本。
    重评：当引入 OCR 或其他 PDF 解析方案时。
    """
    from pypdf import PdfReader

    with Path(file_path).open("rb") as f, PdfReader(f) as reader:
        return "\n\n".join([page.extract_text() for page in reader.pages])


async def parse_pdf_to_text_async(file_path: str) -> str:
    """解析 PDF 并抽取文本（异步，支持对象存储）。

    契约：支持 `s3` 与本地路径；返回按页面拼接的字符串。
    关键路径：读取字节 -> `BytesIO` -> `PdfReader` 解析。
    决策：
    问题：对象存储无法直接提供文件句柄。
    方案：先读取字节并用 `BytesIO` 解析。
    代价：大文件会占用内存。
    重评：当支持流式 PDF 解析时。
    """
    content = await read_file_bytes(file_path)
    with BytesIO(content) as f, PdfReader(f) as reader:
        return "\n\n".join([page.extract_text() for page in reader.pages])


def parse_text_file_to_data(file_path: str, *, silent_errors: bool) -> Data | None:
    """解析文本类文件并返回 `Data`（同步版）。

    契约：支持 `.pdf/.docx/` 及普通文本；失败时 `silent_errors=True` 返回 `None`。
    关键路径：
    1) 根据存储类型选择同步或异步路径；
    2) 按扩展名选择解析器；
    3) 调用 `parse_structured_text` 做结构化解析。
    异常流：解析失败抛 `ValueError`（可静默）。
    决策：
    问题：同步调用方仍需支持对象存储。
    方案：在同步函数内部包装异步路径。
    代价：同步等待可能阻塞线程。
    重评：当上层全面迁移到异步接口时。
    """
    settings = get_settings_service().settings

    # 注意：`s3` 走异步实现并在同步上下文等待结果。
    if settings.storage_type == "s3":
        return run_until_complete(parse_text_file_to_data_async(file_path, silent_errors=silent_errors))

    try:
        if file_path.endswith(".pdf"):
            text = parse_pdf_to_text(file_path)
        elif file_path.endswith(".docx"):
            text = read_docx_file(file_path)
        else:
            text = read_text_file(file_path)

        text = parse_structured_text(text, file_path)
    except Exception as e:
        if not silent_errors:
            msg = f"Error loading file {file_path}: {e}"
            raise ValueError(msg) from e
        return None

    return Data(data={"file_path": file_path, "text": text})


async def parse_text_file_to_data_async(file_path: str, *, silent_errors: bool) -> Data | None:
    """解析文本类文件并返回 `Data`（异步版，支持对象存储）。

    契约：`.pdf` 使用 `BytesIO`；`.docx` 使用临时文件；其他文本直接读取。
    异常流：解析失败抛 `ValueError`（可静默）。
    决策：
    问题：不同格式对读取方式的要求不同。
    方案：按扩展名选择最小依赖的读取策略。
    代价：DOCX 需临时文件，增加 I/O。
    重评：当所有解析器支持字节流输入时。
    """
    try:
        if file_path.endswith(".pdf"):
            text = await parse_pdf_to_text_async(file_path)
        elif file_path.endswith(".docx"):
            text = await read_docx_file_async(file_path)
        else:
            # 实现：普通文本直接读字节并解码，无需临时文件。
            text = await read_text_file_async(file_path)

        # 实现：按扩展名解析结构化文本。
        text = parse_structured_text(text, file_path)

        return Data(data={"file_path": file_path, "text": text})

    except Exception as e:
        if not silent_errors:
            msg = f"Error loading file {file_path}: {e}"
            raise ValueError(msg) from e
        return None


# 注意：暂时移除 `unstructured` 依赖，等待 Python 3.12 支持。
# def get_elements(
#     file_paths: List[str],
#     silent_errors: bool,
#     max_concurrency: int,
#     use_multithreading: bool,
# ) -> List[Optional[Data]]:
#     if use_multithreading:
#         data = parallel_load_data(file_paths, silent_errors, max_concurrency)
#     else:
#         data = [partition_file_to_data(file_path, silent_errors) for file_path in file_paths]
#     data = list(filter(None, data))
#     return data


def parallel_load_data(
    file_paths: list[str],
    *,
    silent_errors: bool,
    max_concurrency: int,
    load_function: Callable = parse_text_file_to_data,
) -> list[Data | None]:
    """并发加载多个文件并返回 `Data` 列表。

    契约：使用线程池并发；输出顺序与输入顺序一致。
    关键路径：创建线程池 -> 并发执行 `load_function` -> 汇总列表。
    决策：
    问题：大量文件串行解析耗时过长。
    方案：线程池并发执行 `load_function`。
    代价：线程数过大可能导致 I/O 争用。
    重评：当解析逻辑改为异步或批处理时。
    """
    with futures.ThreadPoolExecutor(max_workers=max_concurrency) as executor:
        loaded_files = executor.map(
            lambda file_path: load_function(file_path, silent_errors=silent_errors),
            file_paths,
        )
    # 注意：`executor.map` 返回迭代器，需转换为列表以复用结果。
    return list(loaded_files)
